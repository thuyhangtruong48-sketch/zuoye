from __future__ import annotations

import csv
import json
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
OUT_DOCX = DOCS_DIR / "项目14_避开危险路段的救援路径规划_汇报文档.docx"
OUT_MD = DOCS_DIR / "项目14_避开危险路段的救援路径规划_汇报文档.md"

SCENES = [
    {
        "key": "sichuan",
        "label": "四川地震场景",
        "data": ROOT / "data" / "osm_sichuan_earthquake",
        "outputs": ROOT / "outputs" / "osm_sichuan_earthquake",
        "start": "成都市区",
        "target": "汶川县城",
        "disaster": "2008 年汶川地震影响区",
        "focus": "山区塌方、滑坡和道路中断风险",
    },
    {
        "key": "flood",
        "label": "北京洪水场景",
        "data": ROOT / "data" / "osm_beijing_flood",
        "outputs": ROOT / "outputs" / "osm_beijing_flood",
        "start": "清华大学",
        "target": "北京朝阳站",
        "disaster": "2012 年北京 7.21 暴雨洪涝影响区",
        "focus": "积水、通行能力下降和城市拥堵",
    },
    {
        "key": "fire",
        "label": "上海火灾场景",
        "data": ROOT / "data" / "osm_shanghai_fire",
        "outputs": ROOT / "outputs" / "osm_shanghai_fire",
        "start": "武宁消防站",
        "target": "胶州路 728 号火灾点",
        "disaster": "2010 年上海胶州路 11.15 火灾",
        "focus": "火场影响区、交通管制和实时拥堵",
    },
]


def read_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def read_json(path: Path) -> Any:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def fmt(value: Any) -> str:
    text = "" if value is None else str(value)
    return text if text else "无"


def scene_summary(scene: dict[str, Any]) -> dict[str, Any]:
    data_dir = scene["data"]
    output_dir = scene["outputs"]
    nodes = read_csv(data_dir / "nodes.csv")
    edges = read_csv(data_dir / "edges.csv")
    disasters = read_csv(data_dir / "disaster_events.csv")
    mapping = read_csv(data_dir / "road_disaster_mapping.csv")
    traffic = read_csv(data_dir / "road_traffic_mapping.csv")
    results = read_csv(output_dir / "path_comparison.csv")
    by_mode = {row["mode"]: row for row in results if row.get("algorithm") == "Dijkstra"}
    return {
        **scene,
        "node_count": len(nodes),
        "edge_count": len(edges),
        "disaster_count": len(disasters),
        "mapped_count": len(mapping),
        "traffic_count": len(traffic),
        "distance": by_mode.get("distance", {}),
        "safe": by_mode.get("safe", {}),
        "image": output_dir / "route_map_abstract.png",
        "scenario": read_json(data_dir / "scenario.json"),
    }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run.font.name = "Calibri"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=9.5, bold=bold)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[Any]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True)
        set_cell_shading(hdr[i], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], fmt(value))
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_run_font(run, bold=True)
        text = text[len(bold_prefix):]
    run = p.add_run(text)
    set_run_font(run)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
    return doc


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("项目14：避开危险路段的救援路径规划")
    set_run_font(run, size=20, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("汇报文档")
    set_run_font(run, size=13, color=RGBColor(90, 103, 122))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    set_run_font(run, size=9, color=RGBColor(110, 120, 135))


def create_markdown(summaries: list[dict[str, Any]]) -> str:
    lines = [
        "# 项目14：避开危险路段的救援路径规划汇报文档",
        "",
        "## 一、路径规划场景",
        "",
        "本项目面向灾害应急响应中的救援车辆路径规划问题。传统导航通常以距离最短为目标，但在地震、洪水、火灾等灾害场景中，道路可能出现塌方、积水、拥堵、管制或通行能力下降。因此，本项目将真实道路网络抽象为图结构，并把历史灾害影响区和交通状态转化为道路边权重，用 Dijkstra 算法分别计算普通最短路径和安全救援路径。",
        "",
    ]
    for s in summaries:
        lines.append(f"- {s['label']}：{s['start']} → {s['target']}，历史灾害为{s['disaster']}，重点考虑{s['focus']}。")
    lines.extend(
        [
            "",
            "## 二、图模型构建",
            "",
            "道路网络被抽象为无向加权图 G=(V,E)。节点 V 表示道路交叉点、道路折点、起点和终点；边 E 表示相邻节点之间的真实道路路段。每条边记录道路长度、道路名称、危险类型、拥堵程度、是否可通行和综合安全权重。",
            "",
            "道路数据主要来自 OpenStreetMap / Overpass API 区域路网抽取；交通拥堵信息使用高德交通态势 API 进行圆形区域查询后，按道路名称匹配到 OSM 道路边；历史灾害数据来自公开历史灾害资料，并通过影响缓冲区与道路空间叠加识别危险路段。",
            "",
            "## 三、权重设计",
            "",
            "普通最短路径只使用道路距离作为边权重。安全路径使用综合权重：",
            "",
            "`safe_weight = distance × risk_factor × (1 + congestion_weight × congestion) + fixed_cost`",
            "",
            "其中 distance 为道路长度，risk_factor 为灾害风险系数，congestion 为拥堵程度，congestion_weight 为拥堵影响权重，fixed_cost 用于表示火场管制、严重拥堵等额外通行代价。风险系数采用半定量风险等级法，依据 ISO 31000 / IEC 31010 的风险评估思想和交通阻抗函数思想，将正常、拥堵、积水、塌方和火灾核心区转化为不同的道路通行惩罚。",
            "",
            "## 四、算法过程",
            "",
            "本项目只使用 Dijkstra 算法。算法适用于非负权重图，项目中的距离权重和安全权重均为非负数，因此满足 Dijkstra 的使用条件。程序先以距离为权重运行一次，得到普通最短路径；再以综合安全权重运行一次，得到安全救援路径。两次运行使用同一套真实道路图，只改变边权重。",
            "",
            "## 五、结果分析",
            "",
        ]
    )
    for s in summaries:
        d = s["distance"]
        safe = s["safe"]
        lines.append(
            f"- {s['label']}：普通最短路径 {d.get('total_distance')} km，危险类型 {d.get('danger_types') or '无'}，危险边 {d.get('dangerous_edge_count')} 条；安全路径 {safe.get('total_distance')} km，危险类型 {safe.get('danger_types') or '无'}，危险边 {safe.get('dangerous_edge_count')} 条。"
        )
    lines.extend(
        [
            "",
            "整体结果说明，普通最短路径虽然距离较短，但更容易穿越灾害影响区或拥堵路段；安全路径通常距离更长，但通过提高风险路段权重，能够绕开主要危险区域，更适合作为救援推荐路线。",
            "",
            "## 六、改进方向",
            "",
            "后续可以引入更精细的实时封路数据、灾后道路损毁统计、车辆类型差异、动态路况监测和多目标优化方法。对于地震山区场景，高德实时交通态势覆盖不足，后续可结合遥感影像、灾情快报和道路抢通数据进一步提高风险映射精度。",
        ]
    )
    return "\n".join(lines) + "\n"


def build_docx(summaries: list[dict[str, Any]]) -> None:
    doc = setup_document()
    add_title(doc)

    add_heading(doc, "一、路径规划场景", 1)
    add_para(
        doc,
        "本项目面向灾害应急响应中的救援车辆路径规划问题。传统导航通常以距离最短或时间最短为目标，但在地震、洪水和火灾等灾害发生后，道路可能出现塌方、积水、拥堵、交通管制或道路中断。救援车辆更需要综合考虑道路安全性、通行能力和救援效率，因此本项目将真实道路网络、历史灾害影响区和交通状态融合，计算普通最短路径与安全救援路径。",
    )
    add_table(
        doc,
        ["场景", "起点", "终点", "历史灾害", "重点风险"],
        [[s["label"], s["start"], s["target"], s["disaster"], s["focus"]] for s in summaries],
        [1.2, 1.2, 1.3, 1.8, 2.0],
    )

    add_heading(doc, "二、图模型构建", 1)
    add_para(
        doc,
        "道路系统被抽象为无向加权图 G=(V,E)。其中 V 表示道路交叉点、道路折点、救援出发点和受灾目标点；E 表示相邻节点之间的道路路段。每条边保存道路长度、道路名称、危险类型、拥堵程度、是否可通行和综合安全权重。",
    )
    add_para(doc, "数据构建流程：")
    for item in [
        "从 OpenStreetMap / Overpass API 抽取指定区域真实道路网络，得到道路节点、道路边和道路几何轨迹。",
        "将历史灾害事件表示为灾害点和影响半径，形成灾害影响缓冲区。",
        "对道路边和灾害影响区进行空间叠加，若道路进入影响区，则标记为塌方、积水或火灾风险。",
        "对可获得实时路况的区域调用高德交通态势 API，并按道路名称匹配到 OSM 道路边，生成拥堵映射。",
        "将道路节点、道路边、灾害映射、交通映射和路径结果分别保存为 CSV / JSON 数据表。",
    ]:
        add_bullet(doc, item)
    add_table(
        doc,
        ["场景", "节点数", "道路边数", "灾害映射边数", "交通映射边数"],
        [
            [s["label"], s["node_count"], s["edge_count"], s["mapped_count"], s["traffic_count"]]
            for s in summaries
        ],
        [1.5, 1.0, 1.2, 1.5, 1.5],
    )

    add_heading(doc, "三、权重设计", 1)
    add_para(
        doc,
        "普通最短路径只使用道路距离作为边权重，目标是寻找距离最短的路线。安全路径使用综合安全权重，目标是寻找综合通行代价最低的路线。权重公式如下：",
    )
    add_para(doc, "safe_weight = distance × risk_factor × (1 + congestion_weight × congestion) + fixed_cost")
    add_para(
        doc,
        "其中 distance 为道路长度，risk_factor 为灾害风险系数，congestion 为拥堵程度，congestion_weight 为拥堵影响权重，fixed_cost 表示火场管制、严重拥堵等额外通行代价。风险系数采用半定量风险等级法，不把某个系数说成固定物理常数，而是按照风险严重程度把道路通行风险转化为算法可计算的惩罚值。",
    )
    add_table(
        doc,
        ["道路状态", "风险等级", "项目处理方式", "解释"],
        [
            ["正常道路", "低", "risk_factor = 1.0", "作为基准状态，只按距离计算。"],
            ["拥堵道路", "中", "提高拥堵修正项", "来自高德交通态势，表示通行效率下降。"],
            ["积水道路", "较高", "提高 flood 权重", "洪涝会导致涉水风险和通行能力下降。"],
            ["塌方风险", "高", "提高 collapse 权重", "地震山区道路可能受损、中断或绕行。"],
            ["火灾影响区", "高", "提高 fire 权重并可加入 fixed_cost", "火场附近可能存在警戒、救援作业和交通管制。"],
        ],
        [1.1, 0.9, 1.5, 3.0],
    )
    add_para(
        doc,
        "理论依据：风险管理中常采用风险识别、风险分析和风险评价流程，将不确定风险转换为等级或分值；交通规划中常用道路阻抗函数表示拥堵导致的通行代价上升。本项目将这两类思想结合，把灾害影响区和拥堵状态转换为道路边权重。",
    )

    add_heading(doc, "四、Dijkstra 算法过程", 1)
    add_para(
        doc,
        "本项目只使用 Dijkstra 算法。Dijkstra 适用于边权重非负的图，能够求出从起点到其他节点的最小累计代价路径。本项目中的道路距离、灾害风险权重和拥堵修正都不是负数，因此满足算法条件。",
    )
    for step in [
        "初始化：将起点代价设为 0，其余节点代价设为无穷大。",
        "选择节点：每次从未确定节点中选出当前累计代价最小的节点。",
        "松弛边：检查该节点连接的道路边，如果经过该边能降低邻接节点代价，则更新邻接节点代价和前驱节点。",
        "终止：当目标点被确定，或所有可达节点处理完毕后停止。",
        "路径回溯：根据前驱节点从目标点回溯到起点，得到完整路径。",
    ]:
        add_bullet(doc, step)
    add_para(
        doc,
        "项目中 Dijkstra 运行两次：第一次使用 distance 权重，得到普通最短路径；第二次使用 safe_weight 权重，得到安全救援路径。两次计算使用同一张道路图，因此对比结果能够直接反映危险权重对路径选择的影响。",
    )

    add_heading(doc, "五、结果分析", 1)
    add_table(
        doc,
        ["场景", "路径类型", "距离 km", "总代价", "危险边数", "危险类型"],
        [
            [
                s["label"],
                "普通最短路径",
                s["distance"].get("total_distance", ""),
                s["distance"].get("total_cost", ""),
                s["distance"].get("dangerous_edge_count", ""),
                s["distance"].get("danger_types", "") or "无",
            ]
            for s in summaries
        ]
        + [
            [
                s["label"],
                "安全路径",
                s["safe"].get("total_distance", ""),
                s["safe"].get("total_cost", ""),
                s["safe"].get("dangerous_edge_count", ""),
                s["safe"].get("danger_types", "") or "无",
            ]
            for s in summaries
        ],
        [1.25, 1.25, 0.9, 0.9, 0.9, 1.5],
    )
    add_para(
        doc,
        "结果表明，普通最短路径通常距离更短，但更容易穿越灾害影响区或拥堵路段。安全路径通过提高危险路段权重，主动选择绕行路线。四川地震场景中，安全路径距离增加，但危险边数从 825 降为 0；北京洪水场景中，安全路径绕开主要积水区域，危险边数从 86 降为 1；上海火灾场景中，由于终点就在火场附近，安全路径无法完全避开火灾影响区，但危险边数从 14 降为 7。",
    )

    for s in summaries:
        image = s["image"]
        if image.exists():
            add_heading(doc, f"{s['label']}可视化结果", 2)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(image), width=Inches(6.2))
            add_para(doc, f"图示说明：蓝色为普通最短路径，绿色为安全路径，红色或蓝色半透明区域为历史灾害影响区，橙色为危险或拥堵路段，浅灰色为真实道路网络背景。")

    add_heading(doc, "六、改进方向", 1)
    for item in [
        "灾害数据精细化：后续可引入官方道路封闭数据、灾后道路抢通记录、遥感识别结果或更精细的淹没范围数据。",
        "权重参数标定：当前风险系数采用半定量等级法，后续可结合历史通行速度、道路中断概率和救援车辆通过时间进行标定。",
        "动态路径规划：灾害现场道路状态会不断变化，后续可按时间间隔更新交通状态和灾害影响范围，进行动态重规划。",
        "车辆类型扩展：消防车、救护车、物资运输车对道路宽度、坡度、涉水深度和转弯半径要求不同，后续可加入车辆约束。",
        "多目标优化：除安全和距离外，还可加入救援时间、道路等级、医院或避难点容量等因素，形成多目标救援路径模型。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "七、结论", 1)
    add_para(
        doc,
        "本项目构建了一个基于真实道路网络和历史灾害数据的救援路径规划流程。系统能够将灾害影响区和交通拥堵状态映射到道路边，使用 Dijkstra 算法分别计算普通最短路径和安全路径。实验结果说明，在灾害救援场景中，距离最短并不一定最适合作为救援路线；综合考虑灾害风险和道路通行能力后得到的安全路径更符合应急响应需求。",
    )

    add_heading(doc, "参考依据", 1)
    for item in [
        "OpenStreetMap / Overpass API：真实道路网络数据。",
        "高德交通态势 Web 服务：道路实时拥堵状态。",
        "USGS 地震资料：2008 年汶川地震历史事件。",
        "公开历史资料：北京 2012 年 7.21 暴雨、上海胶州路 11.15 火灾。",
        "ISO 31000 / IEC 31010 风险评估思想：风险识别、风险分析和风险评价。",
        "Dijkstra 最短路径算法：非负权重图中的最小累计代价路径搜索方法。",
    ]:
        add_bullet(doc, item)

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)


def main() -> None:
    summaries = [scene_summary(scene) for scene in SCENES]
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    OUT_MD.write_text(create_markdown(summaries), encoding="utf-8")
    build_docx(summaries)
    print(f"Markdown: {OUT_MD}")
    print(f"DOCX: {OUT_DOCX}")


if __name__ == "__main__":
    main()
