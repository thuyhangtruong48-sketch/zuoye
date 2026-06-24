from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "final_presentation"
MD_PATH = OUT_DIR / "项目14_避开危险路段的救援路径规划_详细汇报手稿_真实历史灾害数据版.md"
DOCX_PATH = OUT_DIR / "项目14_避开危险路段的救援路径规划_详细汇报手稿_真实历史灾害数据版.docx"


def set_style_font(style, font_name: str, size: int, color: str | None = None, bold: bool = False) -> None:
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)


def add_footer(document: Document) -> None:
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.text = "项目14：避开危险路段的救援路径规划 - 详细汇报手稿"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def shade_paragraph(paragraph, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    p_pr.append(shd)


def add_body_paragraph(document: Document, text: str, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(11)
    paragraph.paragraph_format.line_spacing = 1.25


def build_docx() -> None:
    text = MD_PATH.read_text(encoding="utf-8")
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = document.styles
    set_style_font(styles["Normal"], "Microsoft YaHei", 11, "1F2937")
    set_style_font(styles["Heading 1"], "Microsoft YaHei", 18, "102033", True)
    set_style_font(styles["Heading 2"], "Microsoft YaHei", 14, "1F5EFF", True)
    set_style_font(styles["Heading 3"], "Microsoft YaHei", 12, "178A4A", True)
    add_footer(document)

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line[2:])
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x10, 0x20, 0x33)
        elif line.startswith("## "):
            document.add_paragraph(line[3:], style="Heading 1")
        elif line.startswith("【页面重点】"):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x17, 0x8A, 0x4A)
            shade_paragraph(paragraph, "E8F6EE")
        elif line.startswith("【讲稿】"):
            paragraph = document.add_paragraph()
            run = paragraph.add_run("讲稿")
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x5E, 0xFF)
            add_body_paragraph(document, line.replace("【讲稿】", ""))
        elif line[0].isdigit() and ". " in line[:4]:
            add_body_paragraph(document, line, "List Number")
        else:
            add_body_paragraph(document, line)

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
